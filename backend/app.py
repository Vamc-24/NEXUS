from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import sys
from datetime import datetime 
import uuid
from dotenv import load_dotenv

# Add project root to path so we can import ai_module
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables
load_dotenv()

from flask_sqlalchemy import SQLAlchemy
from backend.storage import db, SQLAlchemyStorage
from ai_module.pipeline import run_pipeline

app = Flask(__name__, static_folder='..', static_url_path='/')
CORS(app)

# Database Config
# "postgres://..." needed for modern SQLAlchemy (Render provides "postgres://")
uri = os.environ.get('DATABASE_URL', 'sqlite:///nexus.db')
if uri.startswith("postgres://"):
    uri = uri.replace("postgres://", "postgresql://", 1)

app.config['SQLALCHEMY_DATABASE_URI'] = uri
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)

# Create Tables (Local Dev)
with app.app_context():
    db.create_all()

storage = SQLAlchemyStorage(db)

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    data = request.json
    if not data or 'text' not in data:
        return jsonify({'error': 'Invalid data'}), 400
    
    # Extract extended metadata
    feedback_data = {
        'text': data.get('text'),
        'category': data.get('category', 'general'),
        'role': data.get('role', 'Anonymous'),
        'is_verified': data.get('is_verified', False),
        'user_id': data.get('user_id', 'Anonymous'),
        'user_name': data.get('user_name', 'Anonymous'),
        'user_name': data.get('user_name', 'Anonymous'),
        'institute_id': data.get('institute_id', 'Default'),
        'timestamp': datetime.now().isoformat()
    }
    
    record = storage.add_feedback(feedback_data)
    
    return jsonify({'status': 'success', 'id': record['id']}), 201

@app.route('/api/process', methods=['POST'])
def trigger_processing():
    """
    Manually trigger the AI pipeline for a specific institute.
    """
    data = request.json or {}
    institute_id = data.get('institute_id')
    
    try:
        # We need to update run_pipeline signature too, but assuming it takes storage and ID, or just storage 
        # and we pass ID to storage methods?
        # Actually pipeline.py usually calls `storage.get_unprocessed_feedback()`.
        # I need to verify pipeline.py. Assuming I can pass kwargs or update it.
        # For now, let's pass institute_id to run_pipeline if it accepts it, OR 
        # if run_pipeline is blackbox, I might need to update it.
        # Let's assume run_pipeline needs update or we pass it as arg.
        
        # Let's check pipeline import. `from ai_module.pipeline import run_pipeline`
        # I'll pass institute_id and check if I need to update pipeline.py next.
        result = run_pipeline(storage, institute_id=institute_id)
        return jsonify({'status': 'success', 'result': result}), 200
    except Exception as e:
        print(f"Error in pipeline: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/results', methods=['GET'])
@app.route('/api/results', methods=['GET'])
def get_results():
    institute_id = request.args.get('institute_id')
    results = storage.get_latest_results(institute_id=institute_id)
    return jsonify(results)

@app.route('/api/stats', methods=['GET'])
def get_stats():
    institute_id = request.args.get('institute_id')
    if not institute_id or institute_id == 'Default':
        return jsonify(storage.get_feedback_stats())
    return jsonify(storage.get_feedback_stats(institute_id))

@app.route('/api/stats/global', methods=['GET'])
def get_global_stats():
    """Returns global platform statistics for the landing page."""
    return jsonify(storage.get_global_stats())

@app.route('/api/feedback/list', methods=['GET'])
def get_all_feedback_list():
    institute_id = request.args.get('institute_id')
    if not institute_id or institute_id == 'Default':
        # Optional: Decide if we want to show ALL feedback if no ID, or restricted. 
        # For now, consistent with stats, return all if no ID provided (or handle as 'Default')
        return jsonify(storage.get_all_feedback())
    return jsonify(storage.get_all_feedback(institute_id))

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    from fpdf import FPDF
    data = request.json
    results = data.get('results', [])
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Nexus: Strategic Solution Report", ln=1, align='C')
    pdf.ln(10)
    
    def safe_text(text):
        if text is None:
            return ""
        s = str(text)
        # Replace Rupee symbol and other potential non-latin-1 chars
        s = s.replace('₹', 'Rs. ').replace('\u20b9', 'Rs. ')
        # Encode to latin-1, replacing errors to avoid crash
        return s.encode('latin-1', 'replace').decode('latin-1')

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt=safe_text("Nexus: Strategic Solution Report"), ln=1, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt=safe_text(f"Generated on: {datetime.now().strftime('%d %B %Y, %H:%M')}"), ln=1, align='C')
    pdf.ln(10)
    
    for idx, cluster in enumerate(results.get('clusters', [])):
        # Cluster Header
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 12, txt=safe_text(f"#{idx+1} {cluster.get('theme')}"), ln=1, fill=True)
        
        pdf.set_font("Arial", 'I', 11)
        pdf.multi_cell(0, 8, txt=safe_text(f"Issue: {cluster.get('problem_statement')}"))
        pdf.ln(2)
        
        # Solutions
        for sol in cluster.get('solutions', []):
            pdf.set_font("Arial", 'B', 12)
            pdf.set_text_color(41, 151, 255) # Blue accent
            pdf.cell(0, 10, txt=safe_text(f"Strategy: {sol.get('solution_title', 'Proposed Solution')}"), ln=1)
            pdf.set_text_color(0, 0, 0) # Reset color
            
            # Cost & Resources
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(40, 8, txt="Estimated Cost:", border=0)
            pdf.set_font("Arial", '', 10)
            pdf.cell(0, 8, txt=safe_text(f"Rs. {sol.get('total_estimated_cost', 'N/A')}"), ln=1)
            
            res = sol.get('resources', {})
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(40, 8, txt="Investment:", border=0)
            pdf.set_font("Arial", '', 10)
            pdf.cell(0, 8, txt=safe_text(res.get('Investment', '-')), ln=1)

            # Steps
            pdf.ln(2)
            pdf.set_font("Arial", 'B', 10)
            pdf.cell(0, 8, txt="Execution Plan:", ln=1)
            pdf.set_font("Arial", '', 10)
            for step in sol.get('steps', []):
                pdf.cell(10, 8, txt="-")
                pdf.multi_cell(0, 8, txt=safe_text(step))
            
            pdf.ln(5)
        pdf.ln(5)

    filename = f"report_{uuid.uuid4()}.pdf"
    
    # Ensure static directory exists
    static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../frontend')
    filepath = os.path.join(static_dir, filename)
    
    try:
        pdf.output(filepath)
    except Exception as e:
        print(f"Error saving PDF: {e}")
        return jsonify({'error': str(e)}), 500
    
    return jsonify({'url': f'/{filename}'})

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    from docx import Document
    data = request.json
    results = data.get('results', [])
    
    document = Document()
    document.add_heading('Nexus: Student Feedback Analysis Report', 0)
    document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    for cluster in results.get('clusters', []):
        document.add_heading(f"Theme: {cluster.get('theme')}", level=1)
        document.add_paragraph(f"Problem: {cluster.get('problem_statement')}")
        
        for idx, sol in enumerate(cluster.get('solutions', [])):
            document.add_heading(f"Solution {idx+1}: {sol.get('solution_title', 'Solution')}", level=2)
            document.add_paragraph(f"Cost: Rs. {sol.get('total_estimated_cost')}")
            
            res = sol.get('resources', {})
            p = document.add_paragraph()
            p.add_run("Investment: ").bold = True
            p.add_run(f"{res.get('Investment')}\n")
            p.add_run("Labor: ").bold = True
            p.add_run(f"{res.get('Labor')}\n")

            document.add_heading("Steps:", level=3)
            for step in sol.get('steps', []):
                document.add_paragraph(step, style='List Bullet')

    filename = f"report_{uuid.uuid4()}.docx"
    static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../frontend')
    filepath = os.path.join(static_dir, filename)
    
    try:
        document.save(filepath)
    except Exception as e:
        print(f"Error saving DOCX: {e}")
        return jsonify({'error': str(e)}), 500
    
    return jsonify({'url': f'/{filename}'})

@app.route('/api/institute/register', methods=['POST'])
def register_institute():
    data = request.json
    if not data or 'name' not in data:
        return jsonify({'error': 'Institute Name Required'}), 400
        
    try:
        institute_id = storage.register_institute(data)
        return jsonify({'status': 'success', 'id': institute_id}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/institute/verify', methods=['POST'])
def verify_institute():
    data = request.json
    institute_id = data.get('id')
    
    if not institute_id:
        return jsonify({'valid': False}), 400
        
    result = storage.verify_institute(institute_id)
    if result.get('valid'):
        return jsonify({'valid': True, 'name': result['name']}), 200
    return jsonify({'valid': False}), 200

@app.route('/api/institute/register-admin', methods=['POST'])
def register_institute_admin():
    data = request.json
    institute_id = data.get('institute_id')
    admin_id = data.get('admin_id')
    password = data.get('password')

    if not institute_id or not admin_id or not password:
        return jsonify({'error': 'Missing fields'}), 400

    success = storage.register_admin(institute_id, admin_id, password)
    
    if success:
        return jsonify({'status': 'success'})
    else:
        return jsonify({'error': 'Institute not found'}), 404

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    data = request.json
    admin_id = data.get('admin_id')
    password = data.get('password')

    if not admin_id or not password:
        return jsonify({'error': 'Missing credentials'}), 400

    # Dynamic Verification via Storage (Firestore/Local)
    result = storage.verify_admin(admin_id, password)
    if result.get('valid'):
        return jsonify(result)
    else:
        return jsonify({'error': 'Invalid Credentials'}), 401

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
